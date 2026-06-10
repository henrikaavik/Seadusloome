"""Render an impact report as an Estonian-styled ``.docx`` file.

The builder is intentionally thin: it accepts a :class:`Draft` and the
raw ``impact_reports`` row tuple (as fetched in
:func:`app.docs.export_handler.export_report`) and writes a document
to ``EXPORT_DIR/<draft_id>-<report_id>.docx``.

Content structure (per spec §10.3):

    1. Cover heading — "Eelnõu mõjuanalüüsi aruanne" + draft title
    2. Metadata block (uploaded-at, generated-at, ontology version)
    3. "Kokkuvõte" — impact score + affected/conflict/gap counts
    4. "Mõjutatud üksused" — table of affected entities
    5. "Konfliktid" — table of detected conflicts
    6. "EL-i õigusaktide vastavus" — table of EU compliance links
    7. "Lüngad" — table of topic-cluster gaps
    8. Footer with page numbers

The docx layout is plain and uses ``python-docx``'s built-in styles so
the binary stays small and the document opens cleanly in LibreOffice,
Google Docs, and Microsoft Word. A themed template document is a
follow-up (§10.3 note in the spec) — once Riigi Tugiteenuste Keskus
provides the official .dotx we wire it via ``Document(template_path)``.

Environment:

    EXPORT_DIR   Root directory for generated exports. Defaults to
                 ``./storage/exports`` in development and
                 ``/var/seadusloome/exports`` otherwise. Unlike
                 ``STORAGE_DIR`` this value is not secret, so a
                 missing env var off-dev silently falls back to the
                 prod default (no ``RuntimeError``).
"""

from __future__ import annotations

import json
import logging
import os
import shutil
import subprocess
from collections.abc import Callable
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.shared import Pt

from app.docs.draft_model import Draft
from app.docs.labels import TYPE_LABELS_ET as _TYPE_LABELS_ET
from app.ontology.relations import legal_phrase
from app.ui.time import format_tallinn

logger = logging.getLogger(__name__)


# Progress callback signature: ``(current, total)``. ``current`` and
# ``total`` are 1-based work-unit counters; the export progress WS pushes
# them straight to the browser as ``{"current": N, "total": M}``. Called
# from inside ``build_impact_report_docx`` after each major section so
# the UI can render a real ``<progress>`` bar instead of an indeterminate
# spinner (#610). ``None`` disables progress reporting entirely (used by
# the existing test suite + any caller that doesn't need WS push).
ProgressCallback = Callable[[int, int], None]


# How often to publish progress mid-table. Larger tables report after
# every Nth row in addition to the per-section checkpoints; smaller
# tables (< _PROGRESS_BATCH rows) just emit once on entry + once on
# exit via the section-level checkpoints. Tunable here so the WS push
# rate stays reasonable even for outlier 200-row reports (~10 frames
# per second worst case).
_PROGRESS_BATCH = 10


def _safe_publish(callback: ProgressCallback | None, current: int, total: int) -> None:
    """Invoke *callback* swallowing any exception.

    The progress channel is best-effort UX glue; a stuck DB write or a
    flaky Postgres pool must never abort the .docx render. The handler
    layer logs the failure already, so we don't re-log here.
    """
    if callback is None:
        return
    try:
        callback(current, total)
    except Exception:
        logger.debug("docx_export: progress callback failed", exc_info=True)


# Column order used by ``app.docs.export_handler.export_report`` when
# SELECTing from ``impact_reports``. Any schema change needs to be
# mirrored here so the tuple unpacking below stays in sync.
_REPORT_COLUMN_INDEX = {
    "id": 0,
    "draft_id": 1,
    "affected_count": 2,
    "conflict_count": 3,
    "gap_count": 4,
    "impact_score": 5,
    "report_data": 6,
    "ontology_version": 7,
    "generated_at": 8,
}


# ---------------------------------------------------------------------------
# EXPORT_DIR resolution (same pattern as STORAGE_DIR in app/storage/encrypted.py)
# ---------------------------------------------------------------------------


def _load_export_dir() -> Path:
    """Return the root export directory with a dev-friendly default.

    Unlike STORAGE_DIR this value is not secret, so we do not raise a
    RuntimeError when unset in prod — the default prod path is simply
    used as a fallback.
    """
    raw = os.environ.get("EXPORT_DIR")
    if raw:
        return Path(raw)
    if os.environ.get("APP_ENV", "development") == "development":
        return Path("./storage/exports").resolve()
    return Path("/var/seadusloome/exports")


def _get_export_dir() -> Path:
    """Re-read ``EXPORT_DIR`` on every call so tests can monkeypatch it."""
    return _load_export_dir()


# Convenience alias for callers that want to read the directory once.
EXPORT_DIR = _load_export_dir()


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _format_timestamp(value: Any) -> str:
    """Render a ``datetime`` consistently across every docx section (Europe/Tallinn)."""
    return format_tallinn(value)


def _parse_report_data(raw: Any) -> dict[str, Any]:
    """Normalise the ``report_data`` JSONB column into a dict.

    psycopg 3 usually hands back a dict for JSONB columns but older
    drivers (or mocks) may return a JSON string. A missing/empty value
    yields an empty dict so the downstream iteration is a no-op rather
    than crashing the export.
    """
    if raw is None:
        return {}
    if isinstance(raw, dict):
        return raw
    if isinstance(raw, (bytes, bytearray)):
        try:
            return json.loads(raw.decode())
        except (TypeError, ValueError, UnicodeDecodeError):
            logger.warning("docx_export: unparseable bytes in report_data")
            return {}
    if isinstance(raw, str):
        try:
            return json.loads(raw)
        except (TypeError, ValueError):
            logger.warning("docx_export: unparseable string in report_data")
            return {}
    logger.warning("docx_export: unexpected type in report_data: %s", type(raw).__name__)
    return {}


def _short_type(uri: str) -> str:
    """Translate a type URI into an Estonian label where possible."""
    if not uri:
        return "—"
    short = uri.rsplit("#", 1)[-1] if "#" in uri else uri.rsplit("/", 1)[-1]
    return _TYPE_LABELS_ET.get(short, short)


def _relation_cell(row: dict[str, Any]) -> str:
    """Render the "Seose liik" cell from a row's ``relation`` field (#790).

    Mirrors :func:`app.docs.report_routes._relation_cell_text` so the
    .docx export and the HTML view show the exact same Estonian
    legal-language phrase ("muudab", "tõlgendab", "viitab",
    "võtab üle direktiivi", "defineerib mõistet", "on harmoneeritud
    aktiga"). Falls back to ``"—"`` for old impact reports persisted
    before C5 (no ``?relation`` projection) and for gap rows (no single
    predicate per cluster).
    """
    relation = str(row.get("relation") or "").strip()
    if not relation:
        return "—"
    phrase = legal_phrase(relation)
    return phrase or "—"


def _is_partial_match_row(row: dict[str, Any]) -> bool:
    """Return True for ``referencesAct`` literal-edge rows (Wave 2 Step 5A).

    Mirror of :func:`app.docs.report_routes._is_partial_match_row` —
    keeps the .docx export consistent with the HTML view. A
    partial-match row's ``uri`` field is a LITERAL act title (not a
    URL), so the DOCX renderer must not include a hyperlink run for
    it. See docs/2026-05-18-bugfix-plan.md Wave 2 Step 5A.
    """
    relation = str(row.get("relation") or "").strip()
    if not relation:
        return False
    return relation.endswith("referencesAct")


# ---------------------------------------------------------------------------
# Section renderers
# ---------------------------------------------------------------------------


def _add_cover(doc: Any, draft: Draft, report_row: tuple) -> None:
    """Write the cover block — title, draft metadata, generated-at."""
    doc.add_heading("Eelnõu mõjuanalüüsi aruanne", level=0)
    doc.add_heading(draft.title, level=1)

    generated_at = report_row[_REPORT_COLUMN_INDEX["generated_at"]]
    ontology_version = report_row[_REPORT_COLUMN_INDEX["ontology_version"]] or "unknown"

    meta_paragraph = doc.add_paragraph()
    meta_paragraph.add_run("Üles laaditud: ").bold = True
    meta_paragraph.add_run(_format_timestamp(draft.created_at))

    meta_paragraph2 = doc.add_paragraph()
    meta_paragraph2.add_run("Aruanne koostatud: ").bold = True
    meta_paragraph2.add_run(_format_timestamp(generated_at))

    meta_paragraph3 = doc.add_paragraph()
    meta_paragraph3.add_run("Ontoloogia versioon: ").bold = True
    meta_paragraph3.add_run(str(ontology_version))


def _add_summary(doc: Any, report_row: tuple) -> None:
    """Write the "Kokkuvõte" section with score + counts."""
    doc.add_heading("Kokkuvõte", level=1)

    score = report_row[_REPORT_COLUMN_INDEX["impact_score"]]
    affected = report_row[_REPORT_COLUMN_INDEX["affected_count"]]
    conflicts = report_row[_REPORT_COLUMN_INDEX["conflict_count"]]
    gaps = report_row[_REPORT_COLUMN_INDEX["gap_count"]]

    score_paragraph = doc.add_paragraph()
    run = score_paragraph.add_run(f"Mõjuskoor: {score}/100")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph(f"Mõjutatud üksuste arv: {affected}")
    doc.add_paragraph(f"Tuvastatud konfliktide arv: {conflicts}")
    doc.add_paragraph(f"Tuvastatud lünkade arv: {gaps}")


def _add_affected_entities(
    doc: Any,
    findings: dict[str, Any],
    *,
    progress_callback: ProgressCallback | None = None,
    progress_base: int = 0,
    progress_total: int = 0,
) -> int:
    """Write the "Mõjutatud üksused" table.

    Returns the number of work units consumed (1 for the heading +
    1 per ``_PROGRESS_BATCH`` rows). The caller passes ``progress_base``
    (= work units already done) and ``progress_total`` (= grand total)
    so per-row publishes carry the right numerator.
    """
    doc.add_heading("Mõjutatud üksused", level=1)

    rows: list[dict[str, Any]] = list(findings.get("affected_entities") or [])
    if not rows:
        doc.add_paragraph("Mõjutatud üksuseid ei tuvastatud.")
        return 1

    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    header = table.rows[0].cells
    # #790 (C5): "Seose liik" leads each row so the lawyer reading the
    # printed report sees the relation type ("muudab", "tõlgendab", …)
    # before the entity. Keeps parity with the HTML view.
    header[0].text = "Seose liik"
    header[1].text = "Tüüp"
    header[2].text = "Nimetus"
    header[3].text = "URI"
    for index, row in enumerate(rows):
        cells = table.add_row().cells
        cells[0].text = _relation_cell(row)
        # Wave 2 Step 5A (docs/2026-05-18-bugfix-plan.md): partial-match
        # rows (``estleg:referencesAct "<title>"`` literal edge) have
        # no rdf:type and no URI — their ``uri`` field carries the
        # literal act title instead. Surface a dedicated Estonian
        # phrasing for the "Tüüp" column and render the title verbatim
        # in the URI column (no fake URL).
        if _is_partial_match_row(row):
            cells[1].text = "Akt (sätet ei leitud)"
            cells[2].text = str(row.get("label", "") or row.get("uri", "") or "—")
            cells[3].text = str(row.get("uri", "") or "—")
        else:
            cells[1].text = _short_type(str(row.get("type", "")))
            cells[2].text = str(row.get("label", "") or "—")
            cells[3].text = str(row.get("uri", "") or "—")
        # Mid-table checkpoint: every Nth row OR the last row. Avoids
        # firing for tiny tables where the section-level publish above
        # already covers the work.
        if (index + 1) % _PROGRESS_BATCH == 0:
            _safe_publish(
                progress_callback,
                progress_base + 1 + ((index + 1) // _PROGRESS_BATCH),
                progress_total,
            )
    extra_units = len(rows) // _PROGRESS_BATCH
    return 1 + extra_units


def _owned_draft_ids_for_export(org_id: Any) -> set[str]:
    """Return the org's owned draft UUIDs for conflict masking (#844).

    Best-effort — any error (or a falsy ``org_id``) yields an empty set,
    which masks every cross-draft conflict row (the safe default). Kept
    next to :func:`_add_conflicts` so the masking dependency stays local
    to the conflict-rendering region.
    """
    if not org_id:
        return set()
    try:
        from app.db import get_connection
        from app.docs.impact.masking import fetch_owned_draft_ids

        with get_connection() as conn:
            return fetch_owned_draft_ids(conn, str(org_id))
    except Exception:  # noqa: BLE001 — export must never break on masking
        logger.warning(
            "docx_export: owned-draft lookup failed for org=%s; masking all cross-draft rows",
            org_id,
            exc_info=True,
        )
        return set()


def _add_conflicts(
    doc: Any,
    findings: dict[str, Any],
    *,
    owned_draft_ids: set[str] | None = None,
    progress_callback: ProgressCallback | None = None,
    progress_base: int = 0,
    progress_total: int = 0,
) -> int:
    """Write the "Konfliktid" table. See ``_add_affected_entities`` for the progress contract.

    #844 data remediation: the stored ``conflicts`` rows are masked before
    rendering so a report persisted before tenant scoping landed cannot
    leak another org's draft URI / title into the exported .docx. Adhoc
    probe rows are dropped; cross-org draft rows keep their shape but their
    identity is replaced with a neutral Estonian placeholder. The owning
    org's draft UUIDs are supplied by the caller (the export function has
    the ``Draft`` row); ``owned_draft_ids=None`` masks every cross-draft
    row (the safe default).
    """
    from app.docs.impact.masking import drop_adhoc_conflict_rows, mask_conflict_rows

    doc.add_heading("Konfliktid", level=1)

    rows: list[dict[str, Any]] = mask_conflict_rows(
        drop_adhoc_conflict_rows(list(findings.get("conflicts") or [])),
        owned_draft_ids=owned_draft_ids or set(),
    )
    if not rows:
        doc.add_paragraph("Konflikte ei tuvastatud.")
        return 1

    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    header = table.rows[0].cells
    header[0].text = "Seose liik"
    header[1].text = "Eelnõu viide"
    header[2].text = "Konflikti üksus"
    header[3].text = "Põhjus"
    for index, row in enumerate(rows):
        cells = table.add_row().cells
        cells[0].text = _relation_cell(row)
        cells[1].text = str(row.get("draft_ref", "") or "—")
        cells[2].text = str(
            row.get("conflicting_label") or row.get("conflicting_entity", "") or "—"
        )
        cells[3].text = str(row.get("reason", "") or "—")
        if (index + 1) % _PROGRESS_BATCH == 0:
            _safe_publish(
                progress_callback,
                progress_base + 1 + ((index + 1) // _PROGRESS_BATCH),
                progress_total,
            )
    extra_units = len(rows) // _PROGRESS_BATCH
    return 1 + extra_units


def _add_eu_compliance(
    doc: Any,
    findings: dict[str, Any],
    *,
    progress_callback: ProgressCallback | None = None,
    progress_base: int = 0,
    progress_total: int = 0,
) -> int:
    """Write the "EL-i õigusaktide vastavus" table.

    See ``_add_affected_entities`` for the progress contract.
    """
    doc.add_heading("EL-i õigusaktide vastavus", level=1)

    rows: list[dict[str, Any]] = list(findings.get("eu_compliance") or [])
    if not rows:
        doc.add_paragraph("EL-i õigusaktide seoseid ei tuvastatud.")
        return 1

    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    header = table.rows[0].cells
    header[0].text = "Seose liik"
    header[1].text = "EL õigusakt"
    header[2].text = "Eesti säte"
    header[3].text = "Staatus"
    for index, row in enumerate(rows):
        cells = table.add_row().cells
        cells[0].text = _relation_cell(row)
        cells[1].text = str(row.get("eu_label") or row.get("eu_act", "") or "—")
        cells[2].text = str(row.get("provision_label") or row.get("estonian_provision", "") or "—")
        cells[3].text = str(row.get("transposition_status", "") or "—")
        if (index + 1) % _PROGRESS_BATCH == 0:
            _safe_publish(
                progress_callback,
                progress_base + 1 + ((index + 1) // _PROGRESS_BATCH),
                progress_total,
            )
    extra_units = len(rows) // _PROGRESS_BATCH
    return 1 + extra_units


def _add_unresolved_eu_refs(doc: Any, findings: dict[str, Any]) -> int:
    """Write the "EL-i kaardistamata viited" warning section (#815).

    Mirrors the renderer's :func:`_unresolved_eu_refs_section` so the
    exported .docx surfaces the same warning the on-page report does:
    the analyzer detected EU references in the draft text but couldn't
    map them against the ontology snapshot. Returns 0 work units when
    there's nothing to render (legacy reports or clean resolution), so
    progress reporting is unaffected. Returns 1 unit when the section
    is rendered.
    """
    raw = findings.get("unresolved_eu_refs") or []
    if not raw:
        return 0

    seen: set[str] = set()
    unique: list[str] = []
    for entry in raw:
        if not isinstance(entry, dict):
            continue
        ref_text = str(entry.get("ref_text") or "").strip()
        if not ref_text or ref_text in seen:
            continue
        seen.add(ref_text)
        unique.append(ref_text)
    if not unique:
        return 0

    doc.add_heading("EL-i kaardistamata viited", level=1)
    count = len(unique)
    # The persisted ``unresolved_eu_refs`` rows can include both
    # canonical CELEX numbers (e.g. ``32016R0679``) and title/acronym
    # mentions (e.g. ``GDPR``) — the extractor accepts both forms — so
    # the copy says "EL viidet" (EU references), not "CELEX-numbrit".
    doc.add_paragraph(
        f"Tuvastasime dokumendis viiteid EL õigusele "
        f"({count} EL viidet), mida ei õnnestunud ontoloogias "
        "kaardistada:"
    )
    # List each ref as its own bullet so they're easy to scan in the
    # printed report. python-docx's "List Bullet" style ships with the
    # default template.
    for ref_text in unique:
        para = doc.add_paragraph(ref_text, style="List Bullet")
        # Render the ref text in monospace via the run-level font
        # attribute so it visually matches the on-page <code> styling
        # (works equally for CELEX shapes and acronyms).
        for run in para.runs:
            run.font.name = "Courier New"
    doc.add_paragraph(
        "Kontrollige käsitsi — kaardistamata aktid ei kajastu mõjuanalüüsi tulemustes."
    )
    return 1


def _add_gaps(
    doc: Any,
    findings: dict[str, Any],
    *,
    progress_callback: ProgressCallback | None = None,
    progress_base: int = 0,
    progress_total: int = 0,
) -> int:
    """Write the "Lüngad" table. See ``_add_affected_entities`` for the progress contract."""
    doc.add_heading("Lüngad", level=1)

    rows: list[dict[str, Any]] = list(findings.get("gaps") or [])
    if not rows:
        doc.add_paragraph("Lünki ei tuvastatud.")
        return 1

    # Gap rows aggregate provisions per topic cluster, so no single
    # relation predicate is meaningful — the column always renders "—".
    # The column is kept for visual parity with the other three impact
    # sections (#790).
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    header = table.rows[0].cells
    header[0].text = "Seose liik"
    header[1].text = "Teemaklaster"
    header[2].text = "Sätete kaetus"
    header[3].text = "Kirjeldus"
    for index, row in enumerate(rows):
        cells = table.add_row().cells
        cells[0].text = _relation_cell(row)
        cells[1].text = str(row.get("topic_cluster_label") or row.get("topic_cluster", "") or "—")
        cells[
            2
        ].text = f"{row.get('referenced_provisions', '0')} / {row.get('total_provisions', '0')}"
        cells[3].text = str(row.get("description", "") or "—")
        if (index + 1) % _PROGRESS_BATCH == 0:
            _safe_publish(
                progress_callback,
                progress_base + 1 + ((index + 1) // _PROGRESS_BATCH),
                progress_total,
            )
    extra_units = len(rows) // _PROGRESS_BATCH
    return 1 + extra_units


def _add_footer_page_numbers(doc: Any) -> None:
    """Attach a centered ``Page X of Y`` footer to every section.

    python-docx does not expose a high-level "page number" field, so we
    inject the OOXML ``PAGE`` and ``NUMPAGES`` fields directly into the
    footer paragraph. The XML is safe because we control every token;
    no user-supplied text is placed inside the ``<w:instrText>`` field.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    for section in doc.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER

        # Prefix
        paragraph.add_run("Lk ")

        # <w:fldChar w:fldCharType="begin"/>
        run1 = paragraph.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        run1._r.append(fld_begin)

        # <w:instrText xml:space="preserve"> PAGE </w:instrText>
        run2 = paragraph.add_run()
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        run2._r.append(instr)

        # <w:fldChar w:fldCharType="end"/>
        run3 = paragraph.add_run()
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run3._r.append(fld_end)

        paragraph.add_run(" / ")

        # NUMPAGES
        run4 = paragraph.add_run()
        fld_begin2 = OxmlElement("w:fldChar")
        fld_begin2.set(qn("w:fldCharType"), "begin")
        run4._r.append(fld_begin2)

        run5 = paragraph.add_run()
        instr2 = OxmlElement("w:instrText")
        instr2.set(qn("xml:space"), "preserve")
        instr2.text = " NUMPAGES "
        run5._r.append(instr2)

        run6 = paragraph.add_run()
        fld_end2 = OxmlElement("w:fldChar")
        fld_end2.set(qn("w:fldCharType"), "end")
        run6._r.append(fld_end2)


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------


def _compute_progress_total(report_data: dict[str, Any]) -> int:
    """Total work units for one ``build_impact_report_docx`` call.

    The fixed cost is 4 sections (cover + summary + footer + save).
    Each of the four detail tables contributes 1 unit for the heading
    + 1 unit per ``_PROGRESS_BATCH`` rows. We compute it up front so
    the ``<progress>`` bar's ``max`` attribute is right from the first
    push (browsers can render an indeterminate bar if ``value > max``,
    which looks broken).

    The #815 "unresolved EU refs" section contributes 1 unit when
    non-empty and 0 otherwise — it's an alert-style block (no per-row
    table) so the cost is constant.
    """
    fixed = 4  # cover, summary, footer, save
    sections = ("affected_entities", "conflicts", "eu_compliance", "gaps")
    table_units = 0
    for key in sections:
        rows = report_data.get(key) or []
        # 1 unit for the section heading + 1 unit per N rows.
        table_units += 1 + (len(rows) // _PROGRESS_BATCH)
    # #815: optional unresolved-EU-refs warning. Detect by the same
    # rule the helper uses (non-empty list of ref_text-bearing dicts).
    unresolved = report_data.get("unresolved_eu_refs") or []
    if any(
        isinstance(entry, dict) and str(entry.get("ref_text") or "").strip()
        for entry in unresolved
    ):
        table_units += 1
    return fixed + table_units


def build_impact_report_docx(
    draft: Draft,
    report_row: tuple,
    *,
    progress_callback: ProgressCallback | None = None,
) -> Path:
    """Render the impact report as an Estonian-styled ``.docx`` file.

    Writes to ``<EXPORT_DIR>/<draft_id>-<report_id>.docx`` using
    ``python-docx``. The parent directory is created on first write so
    fresh deployments do not need a manual ``mkdir`` step.

    Args:
        draft: The owning draft dataclass (used for title + created_at).
        report_row: Raw tuple fetched in the order declared by
            :data:`_REPORT_COLUMN_INDEX`.
        progress_callback: Optional ``(current, total)`` reporter
            invoked after each major section + every
            ``_PROGRESS_BATCH`` table rows. ``None`` (default) disables
            progress reporting entirely so existing callers and tests
            keep their original behaviour.

    Returns:
        Absolute :class:`pathlib.Path` to the generated file.
    """
    export_dir = _get_export_dir()
    export_dir.mkdir(parents=True, exist_ok=True)

    report_id = report_row[_REPORT_COLUMN_INDEX["id"]]
    out_path = export_dir / f"{draft.id}-{report_id}.docx"

    report_data = _parse_report_data(report_row[_REPORT_COLUMN_INDEX["report_data"]])
    total = _compute_progress_total(report_data)

    doc = Document()

    _add_cover(doc, draft, report_row)
    done = 1
    _safe_publish(progress_callback, done, total)

    _add_summary(doc, report_row)
    done += 1
    _safe_publish(progress_callback, done, total)

    consumed = _add_affected_entities(
        doc,
        report_data,
        progress_callback=progress_callback,
        progress_base=done,
        progress_total=total,
    )
    done += consumed
    _safe_publish(progress_callback, done, total)
    # Section break keeps the detail tables from bleeding into the cover
    # on single-page reports. WD_SECTION.NEW_PAGE is a page break; if the
    # tables are short the blank space is acceptable for a legal export.
    doc.add_section(WD_SECTION.CONTINUOUS)

    # #844: resolve the owning org's draft UUIDs so the conflicts table
    # masks any foreign-org draft identity persisted in a pre-fix report.
    # Best-effort — a lookup failure yields an empty set, which masks
    # every cross-draft row (the safe default).
    consumed = _add_conflicts(
        doc,
        report_data,
        owned_draft_ids=_owned_draft_ids_for_export(draft.org_id),
        progress_callback=progress_callback,
        progress_base=done,
        progress_total=total,
    )
    done += consumed
    _safe_publish(progress_callback, done, total)

    consumed = _add_eu_compliance(
        doc,
        report_data,
        progress_callback=progress_callback,
        progress_base=done,
        progress_total=total,
    )
    done += consumed
    _safe_publish(progress_callback, done, total)

    # #815: render the unresolved-EU-refs warning between EU compliance
    # and Lüngad so it sits next to the EU section it qualifies. The
    # helper returns 0 work units when there's nothing to warn about,
    # which matches the zero contribution from ``_compute_progress_total``.
    consumed = _add_unresolved_eu_refs(doc, report_data)
    done += consumed
    if consumed:
        _safe_publish(progress_callback, done, total)

    consumed = _add_gaps(
        doc,
        report_data,
        progress_callback=progress_callback,
        progress_base=done,
        progress_total=total,
    )
    done += consumed
    _safe_publish(progress_callback, done, total)

    _add_footer_page_numbers(doc)
    done += 1
    _safe_publish(progress_callback, done, total)

    doc.save(str(out_path))
    # Final tick — guarantees the WS sees current == total even when
    # rounding from per-N-row publishes leaves a small gap.
    _safe_publish(progress_callback, total, total)

    logger.info(
        "docx_export: wrote impact report docx draft=%s report=%s path=%s",
        draft.id,
        report_id,
        out_path,
    )
    return out_path


# ---------------------------------------------------------------------------
# C6 (#791) — executive summary printout
# ---------------------------------------------------------------------------


def _coerce_int(value: Any, default: int = 0) -> int:
    """Lenient int coercion — empty / None / malformed lands on *default*."""
    if value is None:
        return default
    try:
        return int(value)
    except (TypeError, ValueError):
        return default


def export_executive_summary(
    draft: Draft,
    report_row: tuple,
) -> Path:
    """Render a 1-2 page executive-summary ``.docx`` for *draft* / *report_row*.

    Intended to be attached to the seletuskiri front matter (the full
    impact report remains as appendix via
    :func:`build_impact_report_docx`).

    Content layout (per the C6 plan section, lines 468-472):

      1. Title: "Mõjuanalüüsi kokkuvõte"
      2. Draft title + author (best-effort from :attr:`Draft.user_id`)
        + generation date.
      3. One-page summary block:

         * Affected provisions count
         * Conflicts count
         * Gaps count
         * Sanctions delta numbers (X uut · Y muudetud · Z eemaldatud)
         * Burden score (before / after / delta-pct) — falls back
           gracefully when the report's ``report_data`` is from before
           C6 and the keys are absent.

    Writes to ``EXPORT_DIR/<draft_id>-<report_id>-summary.docx``.

    Args:
        draft: The owning draft dataclass.
        report_row: Raw ``impact_reports`` tuple keyed by
            :data:`_REPORT_COLUMN_INDEX`.

    Returns:
        Absolute :class:`pathlib.Path` to the generated file.
    """
    export_dir = _get_export_dir()
    export_dir.mkdir(parents=True, exist_ok=True)

    report_id = report_row[_REPORT_COLUMN_INDEX["id"]]
    out_path = export_dir / f"{draft.id}-{report_id}-summary.docx"

    report_data = _parse_report_data(report_row[_REPORT_COLUMN_INDEX["report_data"]])

    doc = Document()

    # --- Title block ------------------------------------------------------
    doc.add_heading("Mõjuanalüüsi kokkuvõte", level=0)
    doc.add_heading(draft.title or "Pealkirjastamata eelnõu", level=1)

    meta1 = doc.add_paragraph()
    meta1.add_run("Üles laaditud: ").bold = True
    meta1.add_run(_format_timestamp(draft.created_at))

    generated_at = report_row[_REPORT_COLUMN_INDEX["generated_at"]]
    meta2 = doc.add_paragraph()
    meta2.add_run("Aruanne koostatud: ").bold = True
    meta2.add_run(_format_timestamp(generated_at))

    # Best-effort "author" line (the draft model carries user_id only;
    # the explicit author name lives in the user table and is not always
    # available to the export job). Render the UUID so the printed
    # summary has a non-empty author field; a future iteration can join
    # ``users.full_name`` once the export handler is wired to the user
    # service.
    author = doc.add_paragraph()
    author.add_run("Autor (kasutaja id): ").bold = True
    author.add_run(str(draft.user_id or "—"))

    # --- One-page summary -------------------------------------------------
    doc.add_heading("Kokkuvõte", level=1)

    score = _coerce_int(report_row[_REPORT_COLUMN_INDEX["impact_score"]])
    affected = _coerce_int(report_row[_REPORT_COLUMN_INDEX["affected_count"]])
    conflicts = _coerce_int(report_row[_REPORT_COLUMN_INDEX["conflict_count"]])
    gaps = _coerce_int(report_row[_REPORT_COLUMN_INDEX["gap_count"]])

    score_para = doc.add_paragraph()
    score_run = score_para.add_run(f"Mõjuskoor: {score}/100")
    score_run.bold = True
    score_run.font.size = Pt(14)

    doc.add_paragraph(f"Mõjutatud sätete arv: {affected}")
    doc.add_paragraph(f"Tuvastatud konfliktide arv: {conflicts}")
    doc.add_paragraph(f"Tuvastatud lünkade arv: {gaps}")

    # Sanctions delta (graceful fallback when the key is missing).
    sanctions_delta = report_data.get("sanctions_delta") or {}
    new_s = _coerce_int(sanctions_delta.get("new_count"))
    mod_s = _coerce_int(sanctions_delta.get("modified_count"))
    rem_s = _coerce_int(sanctions_delta.get("removed_count"))
    sanctions_line = f"Sanktsioonide muutus: {new_s} uut · {mod_s} muudetud · {rem_s} eemaldatud"
    doc.add_paragraph(sanctions_line)

    # Burden score (graceful fallback when missing).
    burden_delta = report_data.get("burden_delta") or {}
    before_score = _coerce_int(burden_delta.get("before_score"))
    after_score_raw = burden_delta.get("after_score")
    delta_pct_raw = burden_delta.get("score_delta_pct")
    if after_score_raw is None and delta_pct_raw is None:
        burden_line = f"Koormuse skoor (kehtivas õiguses, mõjutatud sätete üle): {before_score}"
    else:
        after_str = str(_coerce_int(after_score_raw)) if after_score_raw is not None else "—"
        if delta_pct_raw is None:
            delta_str = "—"
        else:
            d = _coerce_int(delta_pct_raw)
            delta_str = f"{d:+d}%"
        burden_line = (
            f"Koormuse skoor: enne {before_score} · pärast {after_str} · muutus {delta_str}"
        )
    doc.add_paragraph(burden_line)

    # Footer page numbers — reuses the same OOXML helper as the full
    # report so the printed page has a consistent "Lk X / Y" footer.
    _add_footer_page_numbers(doc)

    doc.save(str(out_path))
    logger.info(
        "docx_export: wrote executive summary draft=%s report=%s path=%s",
        draft.id,
        report_id,
        out_path,
    )
    return out_path


_SOFFICE_TIMEOUT_SECONDS = 60


def convert_docx_to_pdf(docx_path: Path) -> Path:
    """Convert an existing .docx into a sibling .pdf via headless LibreOffice (#613).

    The .docx remains the single source of truth for content; PDF is a
    pure visual rendering of the same file. This guarantees the two
    formats can never diverge.

    LibreOffice writes the PDF to ``--outdir`` with the same stem as
    the input file. Returns the absolute path to the generated PDF.

    ``HOME=/tmp`` is set in the subprocess env to silence the benign
    ``dconf-CRITICAL`` warning that fires when the container's HOME
    (e.g. ``/nonexistent``) is not writable. The warning does not
    affect headless conversion but pollutes logs.

    Raises:
        FileNotFoundError: ``soffice`` not on PATH (dev environments
            without LibreOffice installed). Callers must surface a
            clear error rather than retry — installing LibreOffice is
            an operator action, not a transient failure.
        subprocess.TimeoutExpired: conversion exceeded
            :data:`_SOFFICE_TIMEOUT_SECONDS`.
        subprocess.CalledProcessError: ``soffice`` exited non-zero;
            the captured stderr is included in the exception's output.
    """
    soffice = shutil.which("soffice")
    if soffice is None:
        raise FileNotFoundError("soffice not on PATH; LibreOffice is required for PDF export")
    out_dir = docx_path.parent
    env = {**os.environ, "HOME": "/tmp"}
    subprocess.run(
        [
            soffice,
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(out_dir),
            str(docx_path),
        ],
        env=env,
        check=True,
        timeout=_SOFFICE_TIMEOUT_SECONDS,
        capture_output=True,
    )
    pdf_path = out_dir / (docx_path.stem + ".pdf")
    if not pdf_path.exists():
        raise FileNotFoundError(f"LibreOffice did not produce expected PDF at {pdf_path}")
    logger.info(
        "docx_export: converted to pdf docx=%s pdf=%s",
        docx_path,
        pdf_path,
    )
    return pdf_path
