"""Export a chat conversation to Markdown or a .docx file.

Used by the chat routes to give users a downloadable transcript of a
conversation. System messages are skipped — they're internal
orchestration scaffolding, not part of the user-visible dialogue.

Both exporters walk the same ``messages`` iterable and apply a simple
role-to-label mapping:

    user      -> "Kasutaja:"
    assistant -> "Assistent:"
    tool      -> "Tööriist <tool_name>:"

Tool messages render their ``tool_output`` (or ``tool_input`` if no
output is recorded) as pretty-printed JSON inside a fenced code block
(markdown) or as a monospaced paragraph (docx).

Styling for the .docx mirrors the conventions in
``app/drafter/docx_builder.py`` and ``app/docs/docx_export.py`` — Heading
style for the title, normal paragraphs for metadata and body text, bold
inline runs for role labels.
"""

from __future__ import annotations

import io
import json
import logging
import re
from collections.abc import Iterable
from typing import Any

from docx import Document
from docx.shared import Pt

from app.chat.models import Conversation, Message
from app.ui.time import format_tallinn

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Shared helpers
# ---------------------------------------------------------------------------

_ROLE_LABEL_MD = {
    "user": "**Kasutaja:**",
    "assistant": "**Assistent:**",
}

_ROLE_LABEL_PLAIN = {
    "user": "Kasutaja:",
    "assistant": "Assistent:",
}


def _tool_payload(message: Message) -> Any:
    """Return the most useful payload for a tool message.

    Prefers ``tool_output`` (the result) and falls back to ``tool_input``
    (the arguments) when no output is recorded. Returning ``None`` means
    the renderer should emit an empty object placeholder.
    """
    if message.tool_output is not None:
        return message.tool_output
    if message.tool_input is not None:
        return message.tool_input
    return None


def _format_tool_json(payload: Any) -> str:
    """Pretty-print the tool payload as JSON, tolerating non-JSON values."""
    if payload is None:
        return "{}"
    try:
        return json.dumps(payload, ensure_ascii=False, indent=2, sort_keys=True)
    except (TypeError, ValueError):
        logger.exception("Failed to JSON-encode tool payload")
        return str(payload)


# ---------------------------------------------------------------------------
# Markdown export
# ---------------------------------------------------------------------------


def conversation_to_markdown(
    conversation: Conversation,
    messages: Iterable[Message],
) -> str:
    """Render ``conversation`` + ``messages`` as a Markdown string.

    Layout::

        # <title>

        Loodud: DD.MM.YYYY HH:MM

        ---

        **Kasutaja:**

        <content>

        **Assistent:**

        <content>

        **Tööriist <name>:**

        ```json
        { ... }
        ```

    System messages are skipped. Assistant content is emitted verbatim
    (already markdown from the LLM).
    """
    title = (conversation.title or "Vestlus").strip() or "Vestlus"
    parts: list[str] = [
        f"# {title}",
        "",
        f"Loodud: {format_tallinn(conversation.created_at)}",
        "",
        "---",
        "",
    ]

    for message in messages:
        role = message.role
        if role == "system":
            continue

        if role == "tool":
            tool_name = message.tool_name or "tööriist"
            parts.append(f"**Tööriist {tool_name}:**")
            parts.append("")
            parts.append("```json")
            parts.append(_format_tool_json(_tool_payload(message)))
            parts.append("```")
            parts.append("")
            continue

        label = _ROLE_LABEL_MD.get(role)
        if label is None:
            # Unknown role — still render so nothing is lost.
            label = f"**{role}:**"

        parts.append(label)
        parts.append("")
        parts.append(message.content or "")
        parts.append("")

    # Trim trailing blank line, ensure single terminal newline.
    while parts and parts[-1] == "":
        parts.pop()
    return "\n".join(parts) + "\n"


# ---------------------------------------------------------------------------
# .docx export
# ---------------------------------------------------------------------------


# Strip a leading ``#`` block (one or more ``#`` followed by space) from a
# line — converts a markdown heading into flat text without nuking hash
# characters that appear mid-line (e.g. ``#572`` issue references).
_LEADING_HASH_RE = re.compile(r"^#{1,6}\s+")

# Match ``**bold**`` spans. Non-greedy so adjacent bold spans don't merge.
_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")

# Match ``*italic*`` or ``_italic_`` spans (single markers).
_ITALIC_RE = re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")


def _flatten_markdown_line(line: str) -> str:
    """Strip a handful of markdown markers for cleaner docx rendering.

    This is deliberately simple — python-docx has no native markdown
    parser and rolling our own is out of scope. We handle the common
    markers (heading ``#``, bold ``**x**``, italic ``*x*``) and leave
    everything else alone.
    """
    line = _LEADING_HASH_RE.sub("", line)
    line = _BOLD_RE.sub(r"\1", line)
    line = _ITALIC_RE.sub(r"\1", line)
    return line


def _add_role_paragraph(doc: Any, label: str, body: str) -> None:
    """Add a single role-labeled paragraph pair (bold label, then body)."""
    label_para = doc.add_paragraph()
    run = label_para.add_run(label)
    run.bold = True

    if not body:
        return

    for raw_line in body.splitlines() or [body]:
        flat = _flatten_markdown_line(raw_line)
        doc.add_paragraph(flat)


def _add_tool_paragraph(doc: Any, tool_name: str, payload: Any) -> None:
    """Add a tool message as a bold label + monospaced JSON block."""
    label_para = doc.add_paragraph()
    run = label_para.add_run(f"Tööriist {tool_name}:")
    run.bold = True

    json_text = _format_tool_json(payload)
    body_para = doc.add_paragraph()
    body_run = body_para.add_run(json_text)
    body_run.font.name = "Consolas"
    body_run.font.size = Pt(9)


def conversation_to_docx_bytes(
    conversation: Conversation,
    messages: Iterable[Message],
) -> bytes:
    """Render ``conversation`` + ``messages`` as a .docx, returning bytes.

    The returned bytes are a valid ZIP archive (.docx is a zipped OOXML
    bundle) so the caller can stream them directly as a file download.
    """
    title = (conversation.title or "Vestlus").strip() or "Vestlus"

    doc = Document()
    doc.add_heading(title, level=0)
    doc.add_paragraph(f"Loodud: {format_tallinn(conversation.created_at)}")
    doc.add_paragraph("")  # spacer

    for message in messages:
        role = message.role
        if role == "system":
            continue

        if role == "tool":
            _add_tool_paragraph(doc, message.tool_name or "tööriist", _tool_payload(message))
            continue

        label = _ROLE_LABEL_PLAIN.get(role, f"{role}:")
        _add_role_paragraph(doc, label, message.content or "")

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
