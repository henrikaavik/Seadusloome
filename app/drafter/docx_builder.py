"""Build a .docx file from a completed drafter session.

Generates a formatted Estonian legislative document (or VTK memo)
from the session's proposed structure and drafted clauses.

Content structure for full_law:
    1. Cover page: title, date, AI-generated watermark
    2. Table of contents placeholder (Word auto-generates on open)
    3. Main body: chapters / sections / paragraphs
    4. Appendix A: Citation index
    5. Appendix B: Impact analysis summary (if available)
    6. Footer with page numbers

Content structure for VTK:
    1. Cover heading: "VTK eelanaluus" + title
    2. Main body: 5 fixed chapters
    3. Appendix: Citation index
    4. Footer with page numbers
"""

from __future__ import annotations

import logging
import os
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.shared import Pt

logger = logging.getLogger(__name__)


def _get_export_dir() -> Path:
    """Return the export directory, creating it if needed."""
    raw = os.environ.get("EXPORT_DIR")
    if raw:
        p = Path(raw)
    elif os.environ.get("APP_ENV", "development") == "development":
        p = Path("./storage/exports").resolve()
    else:
        p = Path("/var/seadusloome/exports")
    p.mkdir(parents=True, exist_ok=True)
    return p


_WATERMARK_TEXT = (
    "See dokument on genereeritud tehisintellekti abil "
    "Seadusloome AI koostaja kaudu. Palun kontrollige sisu enne kasutamist."
)


def _add_page_number_footer(doc: Any) -> None:
    """Add centered footer with watermark + page number to all sections.

    Combines the AI-generated watermark text (line 1) with the page
    number (line 2) in a single footer so both appear on every page.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    for section in doc.sections:
        footer = section.footer

        # Line 1: watermark text
        watermark_para = footer.paragraphs[0]
        watermark_para.alignment = 1  # CENTER
        wm_run = watermark_para.add_run(_WATERMARK_TEXT)
        wm_run.italic = True
        wm_run.font.size = Pt(8)

        # Line 2: page number
        page_para = footer.add_paragraph()
        page_para.alignment = 1  # CENTER

        page_para.add_run("Lk ")

        run1 = page_para.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        run1._r.append(fld_begin)

        run2 = page_para.add_run()
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        run2._r.append(instr)

        run3 = page_para.add_run()
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run3._r.append(fld_end)

        page_para.add_run(" / ")

        run4 = page_para.add_run()
        fld_begin2 = OxmlElement("w:fldChar")
        fld_begin2.set(qn("w:fldCharType"), "begin")
        run4._r.append(fld_begin2)

        run5 = page_para.add_run()
        instr2 = OxmlElement("w:instrText")
        instr2.set(qn("xml:space"), "preserve")
        instr2.text = " NUMPAGES "
        run5._r.append(instr2)

        run6 = page_para.add_run()
        fld_end2 = OxmlElement("w:fldChar")
        fld_end2.set(qn("w:fldCharType"), "end")
        run6._r.append(fld_end2)


def build_drafter_docx(
    session_id: str,
    title: str,
    workflow_type: str,
    structure: dict[str, Any],
    clauses: list[dict[str, Any]],
    *,
    impact_summary: dict[str, Any] | None = None,
) -> Path:
    """Build the final .docx and return the path.

    Args:
        session_id: UUID string of the drafting session.
        title: Law/VTK title.
        workflow_type: 'full_law' or 'vtk'.
        structure: The proposed structure dict.
        clauses: List of drafted clause dicts.
        impact_summary: Optional impact analysis summary for appendix.

    Returns:
        Path to the generated .docx file.
    """
    export_dir = _get_export_dir()
    out_path = export_dir / f"drafter-{session_id}.docx"

    doc = Document()

    # Build a clause lookup by chapter+paragraph
    clause_map: dict[str, dict[str, Any]] = {}
    for clause in clauses:
        key = f"{clause.get('chapter', '')}/{clause.get('paragraph', '')}"
        clause_map[key] = clause

    if workflow_type == "vtk":
        _build_vtk_doc(doc, title, structure, clause_map)
    else:
        _build_law_doc(doc, title, structure, clause_map)

    # Appendix A: Citation index
    all_citations: list[str] = []
    for clause in clauses:
        all_citations.extend(clause.get("citations", []))

    if all_citations:
        doc.add_section(WD_SECTION.NEW_PAGE)
        doc.add_heading("Lisa A: Viidete register", level=1)
        unique_citations = sorted(set(all_citations))
        for i, cit in enumerate(unique_citations, 1):
            doc.add_paragraph(f"{i}. {cit}")

    # Appendix B: Impact summary (if available)
    if impact_summary:
        doc.add_section(WD_SECTION.NEW_PAGE)
        doc.add_heading("Lisa B: Mojuanaluusi kokkuvote", level=1)

        score = impact_summary.get("impact_score", "N/A")
        doc.add_paragraph(f"Mojuskoor: {score}/100")
        doc.add_paragraph(f"Mojutatud uksuste arv: {impact_summary.get('affected_count', 0)}")
        doc.add_paragraph(f"Tuvastatud konfliktide arv: {impact_summary.get('conflict_count', 0)}")
        doc.add_paragraph(f"Tuvastatud lunkade arv: {impact_summary.get('gap_count', 0)}")

    _add_page_number_footer(doc)

    doc.save(str(out_path))
    logger.info("drafter docx_builder: wrote %s", out_path)
    return out_path


def _build_law_doc(
    doc: Any,
    title: str,
    structure: dict[str, Any],
    clause_map: dict[str, dict[str, Any]],
) -> None:
    """Build the body for a full_law document."""
    # Cover
    doc.add_heading(title or "Eelnou", level=0)
    doc.add_paragraph(f"Koostatud: {datetime.now().strftime('%d.%m.%Y')}")
    doc.add_paragraph("")  # spacer

    # Main body
    for chapter in structure.get("chapters", []):
        chapter_num = chapter.get("number", "")
        chapter_title = chapter.get("title", "")
        doc.add_heading(f"{chapter_num}. {chapter_title}", level=1)

        for section in chapter.get("sections", []):
            para = section.get("paragraph", "")
            section_title = section.get("title", "")
            doc.add_heading(f"{para} {section_title}", level=2)

            key = f"{chapter_num}/{para}"
            clause = clause_map.get(key)
            if clause:
                text = clause.get("text", "")
                if text:
                    doc.add_paragraph(text)
                notes = clause.get("notes", "")
                if notes:
                    p = doc.add_paragraph()
                    run = p.add_run(f"[Markus: {notes}]")
                    run.italic = True
                    run.font.size = Pt(9)
            else:
                doc.add_paragraph("(Sisu puudub)", style="Intense Quote")


def _build_vtk_doc(
    doc: Any,
    title: str,
    structure: dict[str, Any],
    clause_map: dict[str, dict[str, Any]],
) -> None:
    """Build the body for a VTK memo document."""
    # Cover
    doc.add_heading("Vabariigi Valitsuse korralduse eelanaluus", level=0)
    doc.add_heading(title or "VTK eelanaluus", level=1)
    doc.add_paragraph(f"Koostatud: {datetime.now().strftime('%d.%m.%Y')}")
    doc.add_paragraph("")

    # Main body
    for chapter in structure.get("chapters", []):
        chapter_num = chapter.get("number", "")
        chapter_title = chapter.get("title", "")
        doc.add_heading(f"{chapter_num}. {chapter_title}", level=1)

        for section in chapter.get("sections", []):
            para = section.get("paragraph", "")
            section_title = section.get("title", "")
            doc.add_heading(f"{para} {section_title}", level=2)

            key = f"{chapter_num}/{para}"
            clause = clause_map.get(key)
            if clause:
                text = clause.get("text", "")
                if text:
                    doc.add_paragraph(text)
            else:
                doc.add_paragraph("(Sisu puudub)")
