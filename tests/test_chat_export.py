"""Unit tests for ``app.chat.export``.

Exercises both Markdown and .docx renderers against a handcrafted
conversation + message list. No DB access — dataclasses are built
directly.
"""

from __future__ import annotations

import io
import uuid
from datetime import UTC, datetime

from docx import Document

from app.chat.export import (
    conversation_to_docx_bytes,
    conversation_to_markdown,
)
from app.chat.models import Conversation, Message

# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

_CONV_ID = uuid.UUID("33333333-3333-3333-3333-333333333333")
_USER_ID = uuid.UUID("11111111-1111-1111-1111-111111111111")
_ORG_ID = uuid.UUID("22222222-2222-2222-2222-222222222222")


def _make_conversation(title: str = "Testvestlus") -> Conversation:
    created = datetime(2026, 4, 14, 9, 30, tzinfo=UTC)
    return Conversation(
        id=_CONV_ID,
        user_id=_USER_ID,
        org_id=_ORG_ID,
        title=title,
        context_draft_id=None,
        created_at=created,
        updated_at=created,
    )


def _make_message(
    *,
    role: str,
    content: str = "",
    tool_name: str | None = None,
    tool_input: dict | None = None,
    tool_output: dict | None = None,
    offset_minutes: int = 0,
) -> Message:
    return Message(
        id=uuid.uuid4(),
        conversation_id=_CONV_ID,
        role=role,
        content=content,
        tool_name=tool_name,
        tool_input=tool_input,
        tool_output=tool_output,
        rag_context=None,
        tokens_input=None,
        tokens_output=None,
        model=None,
        created_at=datetime(2026, 4, 14, 9, 30 + offset_minutes, tzinfo=UTC),
    )


def _sample_messages() -> list[Message]:
    return [
        _make_message(role="system", content="(sisemine suunis)", offset_minutes=0),
        _make_message(
            role="user",
            content="Kuidas mõjutab see eelnõu äriühingute asutamist?",
            offset_minutes=1,
        ),
        _make_message(
            role="assistant",
            content="# Analüüs\n\n**Peamine mõju** on registreerimiskorra muutus.",
            offset_minutes=2,
        ),
        _make_message(
            role="tool",
            tool_name="sparql_query",
            tool_input={"query": "SELECT ?s WHERE { ?s a ?t }"},
            tool_output={"rows": [{"s": "ex:A"}, {"s": "ex:B"}]},
            offset_minutes=3,
        ),
    ]


# ---------------------------------------------------------------------------
# Markdown
# ---------------------------------------------------------------------------


def test_markdown_contains_title_and_timestamp() -> None:
    conv = _make_conversation(title="Ettevõtlusseaduse analüüs")
    md = conversation_to_markdown(conv, _sample_messages())

    assert md.startswith("# Ettevõtlusseaduse analüüs\n")
    # Europe/Tallinn is UTC+3 in April (DST): 09:30 UTC -> 12:30.
    assert "Loodud: 14.04.2026 12:30" in md


def test_markdown_renders_all_roles_except_system() -> None:
    conv = _make_conversation()
    md = conversation_to_markdown(conv, _sample_messages())

    assert "**Kasutaja:**" in md
    assert "**Assistent:**" in md
    assert "**Tööriist sparql_query:**" in md
    # System message content must be absent.
    assert "sisemine suunis" not in md


def test_markdown_includes_tool_json_fenced_block() -> None:
    conv = _make_conversation()
    md = conversation_to_markdown(conv, _sample_messages())

    assert "```json" in md
    # tool_output preferred over tool_input.
    assert '"rows"' in md
    assert "ex:A" in md
    # tool_input should not appear when tool_output is present.
    assert "SELECT ?s" not in md


def test_markdown_empty_conversation() -> None:
    conv = _make_conversation(title="Tühi vestlus")
    md = conversation_to_markdown(conv, [])

    assert md.startswith("# Tühi vestlus\n")
    assert "**Kasutaja:**" not in md
    assert md.endswith("\n")


def test_markdown_handles_unicode_estonian_chars() -> None:
    conv = _make_conversation(title="Õiguslik küsimus")
    msgs = [
        _make_message(role="user", content="Kas §-s 5 öeldud käive sisaldab käibemaksu?"),
        _make_message(
            role="assistant", content="Jah — § 12 järgi ka siis, kui müüja asub Šveitsis."
        ),
    ]
    md = conversation_to_markdown(conv, msgs)

    assert "Õiguslik küsimus" in md
    assert "§-s 5 öeldud käive" in md
    assert "Šveitsis" in md


def test_markdown_tool_message_without_payload_uses_empty_object() -> None:
    conv = _make_conversation()
    msgs = [_make_message(role="tool", tool_name="noop")]
    md = conversation_to_markdown(conv, msgs)

    assert "**Tööriist noop:**" in md
    assert "```json\n{}\n```" in md


def test_markdown_tool_message_falls_back_to_tool_input() -> None:
    conv = _make_conversation()
    msgs = [
        _make_message(
            role="tool",
            tool_name="sparql_query",
            tool_input={"query": "ASK { ?s ?p ?o }"},
        ),
    ]
    md = conversation_to_markdown(conv, msgs)

    assert "ASK { ?s ?p ?o }" in md


# ---------------------------------------------------------------------------
# .docx
# ---------------------------------------------------------------------------


def test_docx_bytes_are_zip_and_parseable() -> None:
    conv = _make_conversation(title="Dokument")
    data = conversation_to_docx_bytes(conv, _sample_messages())

    assert isinstance(data, bytes)
    assert data[:2] == b"PK"  # ZIP magic bytes.

    # python-docx must be able to reopen the exported blob.
    doc = Document(io.BytesIO(data))
    all_text = "\n".join(p.text for p in doc.paragraphs)

    assert "Dokument" in all_text
    assert "Loodud: 14.04.2026 12:30" in all_text
    assert "Kasutaja:" in all_text
    assert "Assistent:" in all_text
    assert "Tööriist sparql_query:" in all_text
    # System-role content must not leak into the docx.
    assert "sisemine suunis" not in all_text


def test_docx_strips_markdown_markers_in_body() -> None:
    conv = _make_conversation()
    msgs = [
        _make_message(
            role="assistant",
            content="# Pealkiri\n\n**oluline** ja *kaldkirjas* tekst.",
        ),
    ]
    data = conversation_to_docx_bytes(conv, msgs)
    doc = Document(io.BytesIO(data))
    body_text = "\n".join(p.text for p in doc.paragraphs)

    # Heading marker dropped.
    assert "Pealkiri" in body_text
    assert "# Pealkiri" not in body_text
    # Bold/italic markers dropped.
    assert "oluline" in body_text
    assert "kaldkirjas" in body_text
    assert "**oluline**" not in body_text
    assert "*kaldkirjas*" not in body_text


def test_docx_tool_payload_rendered_as_json() -> None:
    conv = _make_conversation()
    msgs = [
        _make_message(
            role="tool",
            tool_name="sparql_query",
            tool_output={"rows": [{"s": "ex:A"}]},
        ),
    ]
    data = conversation_to_docx_bytes(conv, msgs)
    doc = Document(io.BytesIO(data))
    text = "\n".join(p.text for p in doc.paragraphs)

    assert '"rows"' in text
    assert "ex:A" in text


def test_docx_empty_conversation_round_trips() -> None:
    conv = _make_conversation(title="Tühi")
    data = conversation_to_docx_bytes(conv, [])

    assert data[:2] == b"PK"
    doc = Document(io.BytesIO(data))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Tühi" in text


def test_docx_handles_unicode_estonian_chars() -> None:
    conv = _make_conversation(title="Õiguslik küsimus")
    msgs = [
        _make_message(role="user", content="Kas §-s 5 öeldud käive sisaldab käibemaksu?"),
        _make_message(role="assistant", content="Jah — § 12 järgi Šveitsis."),
    ]
    data = conversation_to_docx_bytes(conv, msgs)
    doc = Document(io.BytesIO(data))
    text = "\n".join(p.text for p in doc.paragraphs)

    assert "Õiguslik küsimus" in text
    assert "käibemaksu" in text
    assert "Šveitsis" in text
