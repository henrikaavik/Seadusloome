"""Tests for C6 (#791) — impact-report sanctions + burden delta + executive summary.

Covers:

1. :func:`app.docs.impact.analyzer.analyze_sanctions_delta` —
   per-provision Sanction aggregation, dedup, empty / failure paths.
2. :func:`app.docs.impact.analyzer.analyze_burden_delta` — wraps the
   A2 burden delta and adds before-score + dataclass-friendly fields.
3. :func:`app.docs.docx_export.export_executive_summary` — the 1-2
   page .docx printout has every required field (title, author,
   counts, sanctions numbers, burden score).
4. Legacy reports without ``sanctions_delta`` / ``burden_delta`` keys
   still render: the renderer returns ``""`` and the docx export
   falls back gracefully.
5. Print stylesheet — the ``@media print`` rule forces collapsibles
   open so a browser-print of the report page shows complete content.
"""

from __future__ import annotations

import json
import uuid
from datetime import UTC, datetime
from pathlib import Path
from typing import Any
from unittest.mock import MagicMock

import pytest
from fasthtml.common import to_xml

from app.docs.docx_export import (
    _REPORT_COLUMN_INDEX,
    export_executive_summary,
)
from app.docs.draft_model import Draft
from app.docs.impact.analyzer import (
    BurdenDeltaReport,
    SanctionsDelta,
    SanctionsDeltaRow,
    analyze_burden_delta,
    analyze_sanctions_delta,
)
from app.docs.report_routes import (
    _burden_delta_section,
    _burden_delta_summary_line,
    _print_stylesheet,
    _sanctions_delta_section,
    _sanctions_delta_summary_line,
)

# ---------------------------------------------------------------------------
# Shared fixtures
# ---------------------------------------------------------------------------

_DRAFT_ID = uuid.UUID("aaaaaaaa-aaaa-aaaa-aaaa-aaaaaaaaaaaa")
_REPORT_ID = uuid.UUID("bbbbbbbb-bbbb-bbbb-bbbb-bbbbbbbbbbbb")
_USER_ID = uuid.UUID("cccccccc-cccc-cccc-cccc-cccccccccccc")
_ORG_ID = uuid.UUID("dddddddd-dddd-dddd-dddd-dddddddddddd")

_NS = "https://data.riik.ee/ontology/estleg#"
_KARS_URI = f"{_NS}karistusseadustik"
_KARS_P211_URI = f"{_NS}KarS-p211"
_KARS_P211_SANCTION_URI = f"{_NS}KarS-p211-Sanction"
_OBLIGATION_URI = f"{_NS}NormType_Obligation"
_PROHIBITION_URI = f"{_NS}NormType_Prohibition"


def _make_draft() -> Draft:
    now = datetime.now(UTC)
    return Draft(
        id=_DRAFT_ID,
        user_id=_USER_ID,
        org_id=_ORG_ID,
        title="KarS muutmise eelnõu",
        filename="eelnou.docx",
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        file_size=2048,
        storage_path="/tmp/cipher.enc",
        graph_uri=f"https://data.riik.ee/ontology/estleg/drafts/{_DRAFT_ID}",
        status="ready",
        parsed_text_encrypted=None,
        entity_count=3,
        error_message=None,
        created_at=now,
        updated_at=now,
    )


def _report_row(findings: dict[str, Any]) -> tuple:
    row: list = [None] * len(_REPORT_COLUMN_INDEX)
    row[_REPORT_COLUMN_INDEX["id"]] = _REPORT_ID
    row[_REPORT_COLUMN_INDEX["draft_id"]] = _DRAFT_ID
    row[_REPORT_COLUMN_INDEX["affected_count"]] = len(findings.get("affected_entities") or [])
    row[_REPORT_COLUMN_INDEX["conflict_count"]] = len(findings.get("conflicts") or [])
    row[_REPORT_COLUMN_INDEX["gap_count"]] = len(findings.get("gaps") or [])
    row[_REPORT_COLUMN_INDEX["impact_score"]] = 42
    row[_REPORT_COLUMN_INDEX["report_data"]] = json.dumps(findings)
    row[_REPORT_COLUMN_INDEX["ontology_version"]] = "2026-05-16T00:00+00:00@1"
    row[_REPORT_COLUMN_INDEX["generated_at"]] = datetime(2026, 5, 16, 9, 0, tzinfo=UTC)
    return tuple(row)


def _sanction_sparql_row(
    *,
    provision_uri: str = _KARS_P211_URI,
    provision_label: str = "KarS § 211",
    sanction_uri: str = _KARS_P211_SANCTION_URI,
    sanction_type: str = "imprisonment",
    min_amount: str = "1",
    max_amount: str = "5",
    min_unit: str = "years",
    max_unit: str = "years",
    min_currency: str = "",
    max_currency: str = "",
) -> dict[str, str]:
    return {
        "sanction": sanction_uri,
        "provision": provision_uri,
        "provisionLabel": provision_label,
        "act": _KARS_URI,
        "actLabel": "Karistusseadustik",
        "sanctionType": sanction_type,
        "minAmount": min_amount,
        "maxAmount": max_amount,
        "minUnit": min_unit,
        "maxUnit": max_unit,
        "minCurrency": min_currency,
        "maxCurrency": max_currency,
        "enforcedAtLevel": "act",
        "isStatutoryDefault": "true",
    }


def _burden_sparql_row(
    *,
    provision_uri: str,
    norm_type: str,
    duty_holder: str = "Tööandja",
) -> dict[str, str]:
    return {
        "provision": provision_uri,
        "provisionLabel": f"Säte {provision_uri.rsplit('#', 1)[-1]}",
        "act": _KARS_URI,
        "actLabel": "Karistusseadustik",
        "normType": norm_type,
        "dutyHolder": duty_holder,
    }


# ---------------------------------------------------------------------------
# 1. analyze_sanctions_delta
# ---------------------------------------------------------------------------


class TestAnalyzeSanctionsDelta:
    def test_empty_affected_returns_empty_delta(self) -> None:
        stub = MagicMock()
        delta = analyze_sanctions_delta("urn:draft:1", [], sparql_client=stub)
        assert isinstance(delta, SanctionsDelta)
        assert delta.rows == []
        assert delta.new_count == 0
        assert delta.modified_count == 0
        assert delta.removed_count == 0
        stub.query.assert_not_called()

    def test_blank_uris_are_skipped(self) -> None:
        stub = MagicMock()
        stub.query.return_value = []
        delta = analyze_sanctions_delta("urn:draft:1", ["", "   "], sparql_client=stub)
        assert delta.rows == []
        stub.query.assert_not_called()

    def test_aggregates_one_sanction_per_provision(self) -> None:
        stub = MagicMock()
        # The helper calls list_sanctions_for_provision once per URI.
        # That helper internally calls stub.query — return one Sanction row.
        stub.query.return_value = [_sanction_sparql_row()]
        delta = analyze_sanctions_delta(
            "urn:draft:1",
            [_KARS_P211_URI],
            sparql_client=stub,
        )
        assert len(delta.rows) == 1
        row = delta.rows[0]
        assert row.change == "new"
        assert row.provision_uri == _KARS_P211_URI
        assert row.provision_label == "KarS § 211"
        assert row.sanction_type == "imprisonment"
        assert row.sanction_type_label == "Vangistus"
        assert "1" in row.penalty_range and "5" in row.penalty_range
        assert "aastat" in row.penalty_range
        assert delta.new_count == 1
        assert delta.modified_count == 0
        assert delta.removed_count == 0

    def test_dedup_same_sanction_across_relations(self) -> None:
        """The same provision can surface via several relations in the
        affected-entities pass; dedup keeps one row per (provision, sanction)."""
        stub = MagicMock()
        stub.query.return_value = [_sanction_sparql_row()]
        # Passing the same URI twice simulates the same provision listed
        # multiple times in affected_entities (e.g. via both amends + references).
        delta = analyze_sanctions_delta(
            "urn:draft:1",
            [_KARS_P211_URI, _KARS_P211_URI],
            sparql_client=stub,
        )
        assert len(delta.rows) == 1

    def test_jena_failure_skips_provision(self) -> None:
        stub = MagicMock()
        stub.query.side_effect = RuntimeError("jena down")
        delta = analyze_sanctions_delta(
            "urn:draft:1",
            [_KARS_P211_URI],
            sparql_client=stub,
        )
        # The helper catches the failure inside list_sanctions_for_provision
        # and returns []; the analyze wrapper then produces an empty delta.
        assert delta.rows == []
        assert delta.new_count == 0


# ---------------------------------------------------------------------------
# 2. analyze_burden_delta
# ---------------------------------------------------------------------------


class TestAnalyzeBurdenDelta:
    def test_empty_uri_returns_empty_report(self) -> None:
        stub = MagicMock()
        report = analyze_burden_delta("", sparql_client=stub)
        assert isinstance(report, BurdenDeltaReport)
        assert report.affected_count == 0
        assert report.before_score == 0
        assert report.after_score is None
        assert report.score_delta_pct is None
        # Counts dict fully populated for canonical iteration.
        for key in ("obligation", "prohibition", "permission", "right", "unknown"):
            assert report.counts[key] == 0
        stub.query.assert_not_called()

    def test_aggregates_burden_over_affected_provisions(self) -> None:
        stub = MagicMock()
        # First SPARQL: affected-provisions list.
        # Then one call per provision (the per-provision burden lookup).
        stub.query.side_effect = [
            [{"provision": f"{_NS}P1"}, {"provision": f"{_NS}P2"}],
            [_burden_sparql_row(provision_uri=f"{_NS}P1", norm_type=_OBLIGATION_URI)],
            [_burden_sparql_row(provision_uri=f"{_NS}P2", norm_type=_PROHIBITION_URI)],
        ]
        report = analyze_burden_delta(f"{_NS}Draft_1", sparql_client=stub)
        assert report.affected_count == 2
        assert report.counts["obligation"] == 1
        assert report.counts["prohibition"] == 1
        # before_score = obligations + prohibitions = 2
        assert report.before_score == 2
        # v1: after / pct still None until ontology #214.
        assert report.after_score is None
        assert report.score_delta_pct is None
        # Rows projected with their burden labels.
        assert len(report.rows) == 2
        labels = {r.burden_label for r in report.rows}
        assert "Kohustused" in labels
        assert "Keelud" in labels

    def test_jena_failure_returns_empty_report(self) -> None:
        stub = MagicMock()
        stub.query.side_effect = RuntimeError("jena down")
        report = analyze_burden_delta(f"{_NS}Draft_1", sparql_client=stub)
        assert report.affected_count == 0
        assert report.before_score == 0


# ---------------------------------------------------------------------------
# 3. Renderer — summary lines + section rendering
# ---------------------------------------------------------------------------


class TestSummaryLines:
    def test_sanctions_summary_line_format(self) -> None:
        line = _sanctions_delta_summary_line(
            {"new_count": 3, "modified_count": 1, "removed_count": 0}
        )
        assert "3 uut sanktsiooni" in line
        assert "1 muudetud" in line
        assert "0 eemaldatud" in line

    def test_sanctions_summary_handles_missing_keys(self) -> None:
        line = _sanctions_delta_summary_line({})
        assert "0 uut sanktsiooni" in line
        assert "0 muudetud" in line
        assert "0 eemaldatud" in line

    def test_burden_summary_line_with_delta(self) -> None:
        line = _burden_delta_summary_line(
            {
                "counts": {"obligation": 5, "prohibition": 2, "right": 1, "permission": 0},
                "score_delta_pct": 12,
            }
        )
        assert "5 uut kohustust" in line
        assert "2 keeldu" in line
        assert "1 õigus" in line
        assert "+12%" in line

    def test_burden_summary_line_without_delta(self) -> None:
        line = _burden_delta_summary_line(
            {
                "counts": {"obligation": 5, "prohibition": 2, "right": 1},
                "score_delta_pct": None,
            }
        )
        # Missing / unknown delta falls back to em-dash.
        assert "—" in line

    def test_burden_summary_negative_delta_signed(self) -> None:
        line = _burden_delta_summary_line(
            {
                "counts": {"obligation": 5, "prohibition": 2, "right": 1},
                "score_delta_pct": -8,
            }
        )
        assert "-8%" in line


class TestSectionRendering:
    def _findings_with_delta(self) -> dict[str, Any]:
        return {
            "affected_entities": [],
            "conflicts": [],
            "eu_compliance": [],
            "gaps": [],
            "sanctions_delta": {
                "rows": [
                    {
                        "change": "new",
                        "provision_uri": _KARS_P211_URI,
                        "provision_label": "KarS § 211",
                        "sanction_type": "imprisonment",
                        "sanction_type_label": "Vangistus",
                        "penalty_range": "1 – 5 aastat",
                        "before_summary": "",
                        "after_summary": "Vangistus: 1 – 5 aastat",
                    },
                ],
                "new_count": 1,
                "modified_count": 0,
                "removed_count": 0,
            },
            "burden_delta": {
                "rows": [
                    {
                        "provision_uri": f"{_NS}P1",
                        "provision_label": "Säte P1",
                        "burden_key": "obligation",
                        "burden_label": "Kohustused",
                        "duty_holder": "Tööandja",
                    },
                ],
                "counts": {
                    "obligation": 1,
                    "prohibition": 0,
                    "permission": 0,
                    "right": 0,
                    "unknown": 0,
                },
                "affected_count": 1,
                "before_score": 1,
                "after_score": None,
                "score_delta_pct": None,
            },
        }

    def test_sanctions_section_renders_collapsible_with_summary(self) -> None:
        section = _sanctions_delta_section(self._findings_with_delta(), draft_id=str(_DRAFT_ID))
        xml = to_xml(section)
        # Default-open collapsible.
        assert "<details" in xml
        assert "open" in xml
        # Title + summary line.
        assert "Sanktsioonide muutused" in xml
        assert "1 uut sanktsiooni" in xml
        # Row content.
        assert "KarS § 211" in xml
        assert "Vangistus" in xml
        # Section data attr for the localStorage script.
        assert "sanctions_delta" in xml

    def test_burden_section_renders_with_v1_target_group_label(self) -> None:
        section = _burden_delta_section(self._findings_with_delta(), draft_id=str(_DRAFT_ID))
        xml = to_xml(section)
        assert "<details" in xml
        assert "Koormuse muutused" in xml
        # v1 fallback column label per the plan.
        assert "Kohustatud isik (esialgne, vt #214)" in xml
        assert "Tööandja" in xml
        assert "Kohustused" in xml

    def test_sanctions_section_missing_key_renders_empty(self) -> None:
        # Legacy reports without sanctions_delta render no section at all.
        section = _sanctions_delta_section({}, draft_id=str(_DRAFT_ID))
        assert section == ""

    def test_burden_section_missing_key_renders_empty(self) -> None:
        section = _burden_delta_section({}, draft_id=str(_DRAFT_ID))
        assert section == ""

    def test_sanctions_section_with_empty_rows(self) -> None:
        findings = {
            "sanctions_delta": {
                "rows": [],
                "new_count": 0,
                "modified_count": 0,
                "removed_count": 0,
            }
        }
        section = _sanctions_delta_section(findings, draft_id=str(_DRAFT_ID))
        xml = to_xml(section)
        # Empty-state message shows but the section header still renders.
        assert "ei muuda" in xml or "Sanktsioonide muutused" in xml


class TestPrintStylesheet:
    def test_stylesheet_forces_collapsibles_open_in_print(self) -> None:
        style = _print_stylesheet()
        xml = to_xml(style)
        # The @media print rule is present.
        assert "@media print" in xml
        # And forces a display override on details bodies.
        assert "display: block !important" in xml


# ---------------------------------------------------------------------------
# 4. Executive summary .docx export
# ---------------------------------------------------------------------------


@pytest.fixture
def tmp_export_dir(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> Path:
    monkeypatch.setenv("EXPORT_DIR", str(tmp_path))
    return tmp_path


class TestExecutiveSummaryExport:
    def test_export_writes_a_docx_file(self, tmp_export_dir: Path) -> None:
        findings = {
            "affected_entities": [],
            "conflicts": [],
            "eu_compliance": [],
            "gaps": [],
            "sanctions_delta": {
                "rows": [],
                "new_count": 3,
                "modified_count": 1,
                "removed_count": 0,
            },
            "burden_delta": {
                "rows": [],
                "counts": {
                    "obligation": 5,
                    "prohibition": 2,
                    "permission": 0,
                    "right": 0,
                    "unknown": 0,
                },
                "affected_count": 7,
                "before_score": 7,
                "after_score": None,
                "score_delta_pct": None,
            },
        }
        draft = _make_draft()
        row = _report_row(findings)
        out = export_executive_summary(draft, row)
        assert out.exists()
        assert out.stat().st_size > 0
        # Filename pattern per the spec.
        assert "summary" in out.name
        assert str(_DRAFT_ID) in out.name

    def test_export_contains_required_fields(self, tmp_export_dir: Path) -> None:
        """Verify the rendered .docx text contains every must-have field."""
        from docx import Document as DocxDocument

        findings = {
            "affected_entities": [{"uri": "u1"}] * 4,
            "conflicts": [{"draft_ref": "x"}] * 2,
            "eu_compliance": [],
            "gaps": [{"topic_cluster": "x"}],
            "sanctions_delta": {
                "rows": [],
                "new_count": 3,
                "modified_count": 1,
                "removed_count": 0,
            },
            "burden_delta": {
                "rows": [],
                "counts": {
                    "obligation": 5,
                    "prohibition": 2,
                    "permission": 0,
                    "right": 0,
                    "unknown": 0,
                },
                "affected_count": 7,
                "before_score": 7,
                "after_score": None,
                "score_delta_pct": None,
            },
        }
        draft = _make_draft()
        # The report row's count fields are independently set so they
        # reflect what the lawyer will read on the cover page.
        row_list: list = list(_report_row(findings))
        row_list[_REPORT_COLUMN_INDEX["affected_count"]] = 4
        row_list[_REPORT_COLUMN_INDEX["conflict_count"]] = 2
        row_list[_REPORT_COLUMN_INDEX["gap_count"]] = 1
        row_list[_REPORT_COLUMN_INDEX["impact_score"]] = 42
        out = export_executive_summary(draft, tuple(row_list))

        doc = DocxDocument(str(out))
        text = "\n".join(p.text for p in doc.paragraphs)
        # Title
        assert "Mõjuanalüüsi kokkuvõte" in text
        # Draft title
        assert draft.title in text
        # Author (user id surfaced as the v1 stand-in)
        assert str(_USER_ID) in text
        # Counts
        assert "42" in text  # impact score
        assert "Mõjutatud sätete arv: 4" in text
        assert "Tuvastatud konfliktide arv: 2" in text
        assert "Tuvastatud lünkade arv: 1" in text
        # Sanctions delta numbers.
        assert "3 uut" in text and "1 muudetud" in text and "0 eemaldatud" in text
        # Burden score line.
        assert "Koormuse skoor" in text
        assert "7" in text

    def test_export_works_without_c6_keys(self, tmp_export_dir: Path) -> None:
        """Legacy report (no sanctions_delta / burden_delta) must still export."""
        from docx import Document as DocxDocument

        findings = {
            "affected_entities": [],
            "conflicts": [],
            "eu_compliance": [],
            "gaps": [],
        }
        draft = _make_draft()
        row = _report_row(findings)
        out = export_executive_summary(draft, row)
        assert out.exists()
        # The sanctions / burden lines still render with all-zero / em-dash
        # values rather than NameError.
        doc = DocxDocument(str(out))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Sanktsioonide muutus" in text
        assert "0 uut" in text
        assert "Koormuse skoor" in text


# ---------------------------------------------------------------------------
# 5. Route registration
# ---------------------------------------------------------------------------


class TestRouteRegistration:
    def test_executive_summary_route_is_registered(self) -> None:
        from app.docs.report_routes import register_report_routes

        rt_calls: list[tuple[str, list[str]]] = []

        def _fake_rt(path: str, methods: list[str] | None = None):
            rt_calls.append((path, methods or []))

            def _decorator(fn):
                return fn

            return _decorator

        register_report_routes(_fake_rt)
        paths = {p for p, _ in rt_calls}
        assert "/drafts/{draft_id}/report/summary.docx" in paths


# ---------------------------------------------------------------------------
# 6. analyze_handler integration — JSONB carries the new keys
# ---------------------------------------------------------------------------


class TestAnalyzeHandlerIntegration:
    """The analyze pipeline must persist ``sanctions_delta`` / ``burden_delta``
    keys in ``impact_reports.report_data``."""

    def test_sanctions_delta_dataclass_is_json_friendly(self) -> None:
        """:class:`SanctionsDelta` must round-trip through ``dataclasses.asdict``
        so the analyze handler can serialise it without surprises."""
        import dataclasses

        delta = SanctionsDelta(
            rows=[SanctionsDeltaRow(change="new", provision_uri="u")],
            new_count=1,
            modified_count=0,
            removed_count=0,
        )
        d = dataclasses.asdict(delta)
        # JSON-serialisable via stdlib.
        s = json.dumps(d)
        # Round trips cleanly.
        back = json.loads(s)
        assert back["new_count"] == 1
        assert len(back["rows"]) == 1
        assert back["rows"][0]["change"] == "new"

    def test_burden_delta_dataclass_is_json_friendly(self) -> None:
        import dataclasses

        report = BurdenDeltaReport(
            counts={
                "obligation": 5,
                "prohibition": 2,
                "permission": 0,
                "right": 1,
                "unknown": 0,
            },
            affected_count=8,
            before_score=7,
            after_score=None,
            score_delta_pct=None,
        )
        d = dataclasses.asdict(report)
        s = json.dumps(d)
        back = json.loads(s)
        assert back["before_score"] == 7
        assert back["after_score"] is None
        assert back["counts"]["obligation"] == 5
