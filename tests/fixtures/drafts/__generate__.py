"""Generator script for ``tests/fixtures/drafts/`` sample ``.docx`` files.

Run this once (or whenever the fixture set is updated) to (re)materialise
every ``.docx`` fixture next to this file. The output is deterministic,
small (under a kilobyte each), and committed to the repo so test runs do
not depend on this script firing first.

Usage
-----
::

    uv run python tests/fixtures/drafts/__generate__.py

The script is idempotent — re-running it overwrites the existing files
with the same canonical content. ``python-docx`` is the only third-party
dependency and it is already a project dependency for ``app.docs.docx_export``.

Fixture catalogue
-----------------
``normal_legal_text.docx``   — short Estonian text with 3 §-references and
                                1 EU CELEX reference (happy path).
``very_short.docx``           — title + one short paragraph, no references
                                (tests the empty-extraction code path).
``many_references.docx``      — 15+ §-references, 3 CELEX, 2 court cases
                                (tests extractor recall + dedupe).
``empty_body.docx``           — title only, no body paragraphs (parser
                                edge case).
``malformed_refs.docx``       — typo-ridden references (``§ X.Y.Z``,
                                broken CELEX strings) — tests graceful
                                degradation in the extraction pipeline.

See ``README.md`` in this directory for the per-fixture coverage notes.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.document import Document as _DocumentT

_HERE = Path(__file__).resolve().parent


def _new_doc(title: str) -> _DocumentT:
    """Return a fresh document with *title* applied as ``Heading 1``."""
    doc = Document()
    doc.add_heading(title, level=1)
    return doc


def _write_normal_legal_text() -> Path:
    """Title + 2 paragraphs, 3 §-references + 1 EU CELEX reference."""
    doc = _new_doc("Test eelnõu 1 — tavaline tekst")
    doc.add_paragraph(
        "Käesolev eelnõu täiendab KarS § 5 sätteid ja viitab samuti "
        "§ 12 lg 2 kohustustele. Lisaks tuleb arvestada § 47 "
        "alusel kehtestatud nõuetega."
    )
    doc.add_paragraph(
        "Eelnõu rakendab Euroopa Parlamendi ja nõukogu määruse "
        "32016R0679 (isikuandmete kaitse üldmäärus) põhimõtteid."
    )
    out = _HERE / "normal_legal_text.docx"
    doc.save(out)
    return out


def _write_very_short() -> Path:
    """Title + one short paragraph with no legal references."""
    doc = _new_doc("Test eelnõu 2 — väga lühike")
    doc.add_paragraph("See on lühike eelnõu ilma viideteta.")
    out = _HERE / "very_short.docx"
    doc.save(out)
    return out


def _write_many_references() -> Path:
    """15+ §-references, 3 CELEX strings, 2 court case numbers."""
    doc = _new_doc("Test eelnõu 3 — palju viiteid")

    # 15 §-references across two paragraphs. Mix of bare ``§ N`` and the
    # richer ``§ N lg M`` / ``§ N lg M p K`` shapes so the extractor's
    # regex variants get exercised.
    refs_block_1 = (
        "Eelnõu muudab järgmisi sätteid: § 1, § 2 lg 1, § 3, § 4 lg 2 p 1, "
        "§ 5, § 6, § 7 lg 3, § 8, § 9 lg 1 p 2, § 10."
    )
    refs_block_2 = (
        "Samuti puudutab eelnõu järgmisi sätteid: § 11, § 12 lg 2, "
        "§ 13, § 14 lg 1, § 15, § 16 lg 1 p 3."
    )
    doc.add_paragraph(refs_block_1)
    doc.add_paragraph(refs_block_2)

    # 3 distinct CELEX numbers (regulation + directive + decision).
    doc.add_paragraph(
        "EL õigusaktid: määrus 32016R0679, direktiiv 32019L0790 ja otsus 32020D0001."
    )

    # 2 court-case numbers in the Estonian Supreme Court format.
    doc.add_paragraph(
        "Asjakohane kohtupraktika: Riigikohtu lahend 3-2-1-100-23 ja lahend 3-1-1-50-22."
    )

    out = _HERE / "many_references.docx"
    doc.save(out)
    return out


def _write_empty_body() -> Path:
    """Title only — no body paragraphs."""
    doc = _new_doc("Test eelnõu 4 — tühi sisu")
    out = _HERE / "empty_body.docx"
    doc.save(out)
    return out


def _write_malformed_refs() -> Path:
    """References with typos and broken structure — tests graceful degradation."""
    doc = _new_doc("Test eelnõu 5 — vigased viited")
    doc.add_paragraph(
        "Eelnõu sisaldab vigaseid viiteid: § X.Y.Z (ei ole number), "
        "§ §§ topelt, ning katkist CELEX viidet 320XX0679."
    )
    doc.add_paragraph(
        "Lisaks on pooleli jäänud lahendi number 3-2-1- ja korralik § 5 (mis peaks õnnestuma)."
    )
    out = _HERE / "malformed_refs.docx"
    doc.save(out)
    return out


def generate_all() -> list[Path]:
    """(Re)materialise every fixture file. Returns the list of paths written."""
    return [
        _write_normal_legal_text(),
        _write_very_short(),
        _write_many_references(),
        _write_empty_body(),
        _write_malformed_refs(),
    ]


if __name__ == "__main__":
    written = generate_all()
    for path in written:
        size_kb = path.stat().st_size / 1024
        print(f"  wrote {path.name} ({size_kb:.2f} KiB)")
    print(f"\nGenerated {len(written)} fixture file(s) under {_HERE}.")
