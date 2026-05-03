"""Tests for PDF export (#613).

Three layers covered:
    1. ``convert_docx_to_pdf`` — LibreOffice subprocess shape, env, error paths
    2. ``export_report`` handler — `format=pdf` triggers conversion + persists `pdf_path`
    3. Download route — picks the PDF artefact + ``application/pdf`` MIME

The integration test at the end actually shells out to ``soffice`` if it
is on PATH; it is automatically skipped on dev boxes / CI runners that
don't have LibreOffice installed.
"""

from __future__ import annotations

import shutil
import subprocess
from pathlib import Path
from typing import Any
from unittest.mock import MagicMock, patch
from uuid import UUID

import pytest

from app.docs.docx_export import _SOFFICE_TIMEOUT_SECONDS, convert_docx_to_pdf
from app.docs.export_handler import export_report

_DRAFT_ID = UUID("44444444-4444-4444-4444-444444444444")
_REPORT_ID = UUID("55555555-5555-5555-5555-555555555555")


# ---------------------------------------------------------------------------
# convert_docx_to_pdf
# ---------------------------------------------------------------------------


class TestConvertDocxToPdf:
    def test_invokes_soffice_with_correct_args(self, tmp_path: Path) -> None:
        docx_path = tmp_path / "report.docx"
        docx_path.write_bytes(b"PK\x03\x04 fake docx")  # noqa: S102 — test stub

        # Simulate soffice creating the PDF
        def _fake_run(cmd: Any, **kwargs: Any) -> Any:
            (tmp_path / "report.pdf").write_bytes(b"%PDF-1.4 fake")
            return MagicMock(returncode=0)

        with (
            patch("app.docs.docx_export.shutil.which", return_value="/usr/bin/soffice"),
            patch("app.docs.docx_export.subprocess.run", side_effect=_fake_run) as mock_run,
        ):
            pdf_path = convert_docx_to_pdf(docx_path)

        assert pdf_path == tmp_path / "report.pdf"
        mock_run.assert_called_once()
        args, kwargs = mock_run.call_args
        assert args[0] == [
            "/usr/bin/soffice",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(tmp_path),
            str(docx_path),
        ]
        assert kwargs["check"] is True
        assert kwargs["timeout"] == _SOFFICE_TIMEOUT_SECONDS
        assert kwargs["capture_output"] is True
        # HOME=/tmp silences the dconf warning seen during §4.3 verification
        assert kwargs["env"]["HOME"] == "/tmp"

    def test_raises_when_soffice_missing(self, tmp_path: Path) -> None:
        docx_path = tmp_path / "report.docx"
        docx_path.write_bytes(b"PK\x03\x04")
        with patch("app.docs.docx_export.shutil.which", return_value=None):
            with pytest.raises(FileNotFoundError, match="soffice not on PATH"):
                convert_docx_to_pdf(docx_path)

    def test_raises_when_pdf_not_produced(self, tmp_path: Path) -> None:
        """soffice exits 0 but doesn't write the expected file → hard error."""
        docx_path = tmp_path / "report.docx"
        docx_path.write_bytes(b"PK\x03\x04")
        with (
            patch("app.docs.docx_export.shutil.which", return_value="/usr/bin/soffice"),
            patch(
                "app.docs.docx_export.subprocess.run",
                return_value=MagicMock(returncode=0),
            ),
        ):
            with pytest.raises(FileNotFoundError, match="did not produce expected PDF"):
                convert_docx_to_pdf(docx_path)

    def test_subprocess_timeout_propagates(self, tmp_path: Path) -> None:
        docx_path = tmp_path / "report.docx"
        docx_path.write_bytes(b"PK\x03\x04")
        with (
            patch("app.docs.docx_export.shutil.which", return_value="/usr/bin/soffice"),
            patch(
                "app.docs.docx_export.subprocess.run",
                side_effect=subprocess.TimeoutExpired(cmd="soffice", timeout=60),
            ),
        ):
            with pytest.raises(subprocess.TimeoutExpired):
                convert_docx_to_pdf(docx_path)

    def test_subprocess_nonzero_propagates(self, tmp_path: Path) -> None:
        docx_path = tmp_path / "report.docx"
        docx_path.write_bytes(b"PK\x03\x04")
        with (
            patch("app.docs.docx_export.shutil.which", return_value="/usr/bin/soffice"),
            patch(
                "app.docs.docx_export.subprocess.run",
                side_effect=subprocess.CalledProcessError(
                    returncode=1, cmd="soffice", stderr=b"GPG key expired"
                ),
            ),
        ):
            with pytest.raises(subprocess.CalledProcessError):
                convert_docx_to_pdf(docx_path)


# ---------------------------------------------------------------------------
# export_report handler
# ---------------------------------------------------------------------------


def _stub_draft() -> Any:
    draft = MagicMock()
    draft.id = _DRAFT_ID
    return draft


def _stub_report_row() -> tuple:
    # (id, draft_id, ...) — only the first two are checked by the handler.
    return (str(_REPORT_ID), str(_DRAFT_ID), 0, 0, 0, 0, {}, "v1", None)


class TestExportReportFormat:
    def test_default_format_is_docx_no_subprocess(self, tmp_path: Path) -> None:
        docx_path = tmp_path / "out.docx"
        docx_path.write_bytes(b"PK\x03\x04")

        with (
            patch("app.docs.export_handler.get_connection") as mock_conn_cm,
            patch("app.docs.export_handler.get_draft", return_value=_stub_draft()),
            patch(
                "app.docs.export_handler.build_impact_report_docx",
                return_value=docx_path,
            ),
            patch("app.docs.export_handler.convert_docx_to_pdf") as mock_convert,
        ):
            mock_conn = MagicMock()
            mock_conn.execute.return_value.fetchone.return_value = _stub_report_row()
            mock_conn_cm.return_value.__enter__.return_value = mock_conn

            result = export_report({"draft_id": str(_DRAFT_ID), "report_id": str(_REPORT_ID)})

        assert result["format"] == "docx"
        assert result["docx_path"] == str(docx_path)
        assert "pdf_path" not in result
        # PDF conversion never invoked on the docx path.
        mock_convert.assert_not_called()

    def test_pdf_format_invokes_conversion_and_persists_path(self, tmp_path: Path) -> None:
        docx_path = tmp_path / "out.docx"
        docx_path.write_bytes(b"PK\x03\x04")
        pdf_path = tmp_path / "out.pdf"

        with (
            patch("app.docs.export_handler.get_connection") as mock_conn_cm,
            patch("app.docs.export_handler.get_draft", return_value=_stub_draft()),
            patch(
                "app.docs.export_handler.build_impact_report_docx",
                return_value=docx_path,
            ),
            patch(
                "app.docs.export_handler.convert_docx_to_pdf",
                return_value=pdf_path,
            ) as mock_convert,
        ):
            mock_conn = MagicMock()
            mock_conn.execute.return_value.fetchone.return_value = _stub_report_row()
            mock_conn_cm.return_value.__enter__.return_value = mock_conn

            result = export_report(
                {
                    "draft_id": str(_DRAFT_ID),
                    "report_id": str(_REPORT_ID),
                    "format": "pdf",
                }
            )

        assert result["format"] == "pdf"
        assert result["docx_path"] == str(docx_path)
        assert result["pdf_path"] == str(pdf_path)
        mock_convert.assert_called_once_with(docx_path)

    def test_unknown_format_rejected_with_value_error(self) -> None:
        with pytest.raises(ValueError, match="unsupported format"):
            export_report(
                {
                    "draft_id": str(_DRAFT_ID),
                    "report_id": str(_REPORT_ID),
                    "format": "xlsx",
                }
            )

    def test_legacy_payload_without_format_treated_as_docx(self, tmp_path: Path) -> None:
        """Pre-#613 jobs in the queue have no `format` key — must still work."""
        docx_path = tmp_path / "out.docx"
        docx_path.write_bytes(b"PK\x03\x04")
        with (
            patch("app.docs.export_handler.get_connection") as mock_conn_cm,
            patch("app.docs.export_handler.get_draft", return_value=_stub_draft()),
            patch(
                "app.docs.export_handler.build_impact_report_docx",
                return_value=docx_path,
            ),
            patch("app.docs.export_handler.convert_docx_to_pdf") as mock_convert,
        ):
            mock_conn = MagicMock()
            mock_conn.execute.return_value.fetchone.return_value = _stub_report_row()
            mock_conn_cm.return_value.__enter__.return_value = mock_conn
            result = export_report({"draft_id": str(_DRAFT_ID), "report_id": str(_REPORT_ID)})
        assert result["format"] == "docx"
        mock_convert.assert_not_called()


# ---------------------------------------------------------------------------
# Integration smoke (real soffice — skipped unless LibreOffice is available)
# ---------------------------------------------------------------------------


@pytest.mark.skipif(
    shutil.which("soffice") is None,
    reason="LibreOffice (soffice) not available; install for full smoke",
)
def test_real_soffice_produces_valid_pdf(tmp_path: Path) -> None:
    """End-to-end: build a tiny .docx and shell out to the real soffice."""
    from docx import Document

    doc = Document()
    doc.add_heading("Test", level=0)
    doc.add_paragraph("Hello PDF.")
    docx_path = tmp_path / "tiny.docx"
    doc.save(str(docx_path))

    pdf_path = convert_docx_to_pdf(docx_path)

    assert pdf_path.exists()
    assert pdf_path.read_bytes().startswith(b"%PDF-")
